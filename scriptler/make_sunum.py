from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2); section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
BLUE = RGBColor(0x1A,0x56,0xDB); DARK = RGBColor(0x1E,0x40,0xAF); GRAY = RGBColor(0x6B,0x72,0x80)

def h1(t):
    p=doc.add_heading(t,1); p.runs[0].font.color.rgb=BLUE
def h2(t):
    p=doc.add_heading(t,2); p.runs[0].font.color.rgb=DARK
def para(t,bold=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold
def b(t,px=None):
    p=doc.add_paragraph(style='List Bullet')
    if px: r2=p.add_run(px+': '); r2.bold=True
    p.add_run(t)
def tbl(hdrs,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(hdrs)); t.style='Table Grid'
    hr=t.rows[0]
    for i,h in enumerate(hdrs):
        c=hr.cells[i]; c.text=h
        c.paragraphs[0].runs[0].bold=True
        c.paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tp=tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
        s.set(qn('w:fill'),'1A56DB'); s.set(qn('w:val'),'clear'); tp.append(s)
    for ri,rd in enumerate(rows):
        rw=t.rows[ri+1]
        for ci,ct in enumerate(rd): rw.cells[ci].text=str(ct)

# KAPAK
tt=doc.add_heading('MAOS V5',0); tt.alignment=WD_ALIGN_PARAGRAPH.CENTER
tt.runs[0].font.color.rgb=BLUE; tt.runs[0].font.size=Pt(36)
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('Multi-Agent Orchestration System'); r.font.size=Pt(18); r.bold=True; r.font.color.rgb=DARK
doc.add_paragraph()
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
s2.add_run('Cok Ajanli Yapay Zeka ile Tam Otomatik Yazilim Gelistirme Sistemi')
doc.add_paragraph()
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
m.add_run(f'Ahmet Sayin - {datetime.datetime.now().strftime("%B %Y")}')
doc.add_page_break()

# 1 - OZET
h1('1. YONETICI OZETI')
para('MAOS V5, kullanicinin dogal dille verdigi yazilim gelistirme gorevini 9 ozell ajandan olusan bir takim araciligiyla insan mudahalesi olmadan, tam ve calisir koda donusturuyor.')
doc.add_paragraph()
b('Tek LLM hatalarini cok ajanli sistemle sinirlandirma','Uzmanlaşma')
b('Planlama, Kodlama, Test, Kalite Denetimi tam dongusu','Tam Dongu')
b('Her ajana en uygun, maliyet-verimli model atamasi','Model Yonlendirme')
b('Basarisiz gorevlerin otomatik retry ve hata ogrenmesi','Oz-Iyilestirme')
b('EC2 uzerinde 7/24 Dashboard + Telegram Bot + REST API','Canli Sistem')
doc.add_page_break()

# 2 - MIMARI
h1('2. SISTEM MIMARISI')
h2('2.1. Katmanlar')
b('Web Dashboard, Telegram Bot, REST API, CLI','Erisim')
b('FastAPI + WebSocket/SSE streaming','API')
b('Orchestrator, Message Bus, Memory Manager','Orkestrasyon')
b('9 Uzman Ajan (Planner, Coder, Critic, Executor vb.)','Ajanlar')
b('OpenRouter, Groq, Cerebras LLM saglayicilari','LLM')
b('ChromaDB vektör hafiza + JSON proje hafizasi','Hafiza')
doc.add_paragraph()
h2('2.2. Gorev Akisi')
tbl(['#','Bilesen','Yapilan Is'],[
    ['1','Kullanici','Dogal dille hedef girer (orn: Flask API yaz)'],
    ['2','Orchestrator','Gorevi alir, Planner cagiriyor'],
    ['3','Planner','Alt gorevler + ajan atamalari JSON olarak uretilir'],
    ['4','Researcher','Web aramasi yapar, yaklasim onerisi sunar'],
    ['5','Coder','Kodu [FILE:] bloklari ile dosyalara yazar'],
    ['6','Critic','1-10 puan; 5+ onay, dusuk ise Codeer\'a geri gonder'],
    ['7','Linter','Flake8/Pylint statik analiz (LLM yok, ucretsiz)'],
    ['8','Tester','pytest otomatik test (LLM yok, ucretsiz)'],
    ['9','Executor','Interaktif terminal, runtime hatalarini cozer'],
    ['10','Memory','Basarili projeyi ChromaDB vektör hafizaya yazar'],
    ['11','Kullanici','workspace/ klasorunde hazir, test edilmis kod'],
])
doc.add_page_break()

# 3 - AJANLAR
h1('3. UZMAN AJANLAR')
h2('3.1. Planner — Bas Mimar')
b('Kullanici hedefini alt gorevlere boler, ajan atamalari yapar')
b('Gorev bagimliliklar grafigi (depends_on) olusturur')
b('Ciktiyi kesinlikle JSON formatinda uretir, hic arac kullanmaz')
b('Model: qwen3-30b-a3b (guclu reasoning)')
doc.add_paragraph()
h2('3.2. Researcher — Bilgi Sentezleyici')
b('DuckDuckGo ile web aramasi, guncel en iyi yaklasimi bulur')
b('Benzer proje hafizada varsa aramayi atlar (maliyet tasarrufu)')
b('"Onerilen Yaklasim + Riskler + Kod Iskeleti" formatinda rapor')
b('Model: gpt-oss-120b')
doc.add_paragraph()
h2('3.3. Coder — Isci Ari')
b('[FILE: src/main.py]...[/FILE] bloklari ile dosya uretir')
b('Otomatik sozdizimi onarimi (regex ile eksik kapanislar duzeltilir)')
b('Artirimli duzenleme: sadece degisen satirlari yazar (Git diff tiz)')
b('Model: Codestral-2508 (256K output, kodlama sertifikali)')
doc.add_paragraph()
h2('3.4. Critic — Kalite Denetcisi')
b('5 kriter: dogruluk, kalite, test kapsami, mimari, guvenlik')
b('1-10 puan, esik 5.0; alti ise hata raporu ile Coder\'a iade')
b('Model: step-3.5-flash (hizli, ucuz QA modeli)')
doc.add_paragraph()
h2('3.5. Executor — Interaktif Terminal')
b('Kod terminalde calistirilir, cikti LLM\'e gosterilir')
b('Hata -> Dusun -> Arac Kullan (shell/read/write) -> Coz dongusu')
b('pip install ile eksik paketleri otomatik kurar (max 3 tur)')
doc.add_paragraph()
h2('3.6. Linter & Tester — LLM\'siz Denetciler')
b('SIFIR AI MALIYETI — hic model cagrisi yok')
b('Linter: Flake8 + Pylint (sozdizimi + stil)')
b('Tester: pytest (test_*.py otomatik kesfedilir ve calistirilir)')
doc.add_paragraph()
h2('3.7. Builder — APK Ureticisi')
b('Telegram /build <aciklama> komutuyla tetiklenir')
b('Flet projelerini Android APK\'ya otomatik derler')
b('APK hazir olunca Telegram\'a dosya olarak gonderilir')
doc.add_paragraph()
h2('3.8. Memory Agent — Deneyim Yoneticisi')
b('Basarili projeleri ChromaDB vektör veritabanina yazar')
b('Yeni projede benzer gecmisi arar, Planner\'a rehber sunar')
b('Basarisiz gorevler "kacinilmasi gereken pattern" olarak isaretlenir')
b('Aktif Ogrenme: similarity x (1+success) / (1+cost*5) formulu')
doc.add_page_break()

# 4 - CEKIRDEK
h1('4. CEKIRDEK BILESENLEER')
h2('Orchestrator'); b('ReAct dongusu + Critic-Coder geri bildirim (max 3 tur)')
b('asyncio ile paralel gorev calistirma'); b('Fazli proje yonetimi (Phase 1 ciktisi Phase 2 girissi)')
doc.add_paragraph()
h2('LLM Client'); b('3 saglayici: OpenRouter, Groq, Cerebras')
b('Rate Limiter + Exponential Backoff (2->4->8 sn)'); b('Token maliyeti anlik USD hesabi + SSE streaming')
doc.add_paragraph()
h2('Vector Memory'); b('ChromaDB + sentence-transformers, semantik arama')
b('Basari + maliyet agirlikli siralamasi (Aktif Ogrenme)')
doc.add_paragraph()
h2('Message Bus'); b('Pub/Sub mimarisi, CRITICAL>HIGH>NORMAL>LOW oncelik')
b('Tum mesajlar messages.jsonl\'a loglanir')
doc.add_page_break()

# 5 - MODEL
h1('5. MODEL YONLENDIRME')
tbl(['Ajan','Model','Saglayici','Giris Fiyati','Neden?'],[
    ['Planner','Qwen3-30B-A3B','OpenRouter','$0.07/M','Reasoning, JSON uretimi'],
    ['Researcher','GPT-OSS-120B','OpenRouter','$0.04/M','Bilgi sentezi'],
    ['Coder','Codestral-2508','OpenRouter','$0.30/M','256K output, kodlama'],
    ['Coder Fast','Gemini Flash Lite','OpenRouter','$0.25/M','Hizli gorevler'],
    ['Critic','Step-3.5-Flash','OpenRouter','$0.10/M','Hizli QA'],
    ['Executor','GPT-OSS-120B','OpenRouter','$0.04/M','Terminal reasoning'],
    ['Linter','YOK (Flake8)','-','$0.00','LLM yok, ucretsiz'],
    ['Tester','YOK (pytest)','-','$0.00','LLM yok, ucretsiz'],
])
doc.add_page_break()

# 6 - CANLI
h1('6. CANLIYA ALINAN OZELLIKLER')
h2('Web Dashboard (98.80.144.72:8000/dashboard)')
b('AWS EC2 Ubuntu 22.04, 7/24 systemd servisi')
b('Aktif/gecmis session + gorev loglari'); b('Model bazli USD maliyet istatistikleri')
b('10 sn otomatik yenileme, premium dark theme')
doc.add_paragraph()
h2('Telegram Bot'); b('/start, /status, /build <proje> komutlari')
b('Build bitince APK dosyasi bot mesajina eklenerek gonderilir')
doc.add_paragraph()
h2('REST API')
tbl(['Endpoint','Method','Aciklama'],[
    ['/run','POST','Yeni proje baslt'],
    ['/status/{id}','GET','Session durumu'],
    ['/api/sessions','GET','Tum oturumlar'],
    ['/api/costs','GET','Maliyet'],
    ['/api/tasks','GET','Son gorevler'],
    ['/ws/{id}','WS','Canli log akisi'],
])
doc.add_page_break()

# 7 - TEKNOLOJI
h1('7. TEKNOLOJI YIGINI')
tbl(['Katman','Teknoloji'],[
    ['Backend','Python 3.12, FastAPI, asyncio'],
    ['LLM Iletisim','httpx async HTTP, SSE streaming'],
    ['LLM Saglayicilar','OpenRouter, Groq, Cerebras'],
    ['Vektör Hafiza','ChromaDB + sentence-transformers'],
    ['Uzun Donem Hafiza','JSON + Pydantic modeller'],
    ['Kod Analizi','Python AST modülü + regex'],
    ['Test & Lint','pytest, flake8, pylint'],
    ['Deployment','AWS EC2 Ubuntu 22.04, systemd'],
    ['Bot','python-telegram-bot v20 async'],
    ['Frontend','Tailwind CSS CDN, vanilla JS'],
    ['Sandboxing','Path validation (workspace izolasyonu)'],
])
doc.add_page_break()

# 8 - SONUC
h1('8. SONUC')
para('MAOS V5 cok ajanli koordinasyon, adaptif hafiza ve maliyet optimizasyonu ile ozgun bir arastirma projesidir.')
doc.add_paragraph()
h2('Basarilar')
b('Gercek projeleri uctan uca otomatik uretme (Flask API, CLI, mobil)')
b('Aktif ogrenme ile her projeden deneyim birikimi')
b('AWS EC2 7/24 canli sistem + Telegram Bot + Dashboard')
doc.add_paragraph()
h2('Gelecek Hedefler')
b('E-ticaret ve mobil app gibi buyuk projelerde test')
b('Multi-modal UI tester (ekran goruntusu degerlendirme)')
b('Docker container izolasyonu | TypeScript, Go, Java destegi')

out = r'c:\Users\ahmed\OneDrive\Masaustu\MAOS_V5_Sunum.docx'
try:
    doc.save(out)
    print(f'HAZIR: {out}')
except Exception:
    out = r'c:\Users\ahmed\Desktop\MAOS_V5_Sunum.docx'
    doc.save(out)
    print(f'HAZIR (Desktop): {out}')
