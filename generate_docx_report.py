import subprocess
import sys
import os

try:
    import docx
except ImportError:
    print("python-docx is missing. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo
        else:
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('MAOS (Multi-Agent Orchestration System) V5', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Kapsamlı Sistem Mimarisi ve Analiz Raporu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    # 1. Executive Summary
    add_heading(doc, '1. Yönetici Özeti (Executive Summary)', 1)
    doc.add_paragraph(
        "MAOS V5, karmaşık yazılım geliştirme, analiz ve otomasyon görevlerini otonom "
        "olarak parçalara ayırarak yürüten nap-tabanlı (multi-agent) bir orkestrasyon sistemidir. "
        "ReAct (Reason-Act) metodolojisiyle çalışan sistem; web araması (RAG destekli), kod yazımı, "
        "statik/dinamik testler, kullanıcı arayüzü testleri (Playwright) ve hatta mobil build yeteneklerine sahiptir."
    )
    
    # 2. Mimari Yapı
    add_heading(doc, '2. Çekirdek Mimari (Core Architecture)', 1)
    
    add_heading(doc, '2.1 Orchestrator (Orkestratör)', 2)
    doc.add_paragraph(
        "Sistemin ana yönetim merkezidir. Kullanıcıdan gelen talebi (goal) alır, bunu görev planına dönüştürür "
        "ve ajanlara dağıtır. Hem 'phased' (aşamalı) hem de 'flat' (tekdüze) planları yönetebilir. "
        "Kendi içinde bir görev hafızası (task record) oluşturarak PENDING, IN_PROGRESS, COMPLETED, FAILED "
        "gibi durumları idare eder."
    )
    
    add_heading(doc, '2.2 İletişim & Hafıza (Message Bus & Memory)', 2)
    doc.add_paragraph(
        "Ajanlar arası iletişim pub/sub yapısına sahip asenkron 'Message Bus' üzerinden sağlanır. "
        "Veri kalıcılığı ve bağlam belleği için 'Memory Agent' ile ChromaDB/Vector bellek entegrasyonu mevcuttur. "
        "Bu sayede sistem daha önceki konuşmaları ve geliştirilen projeleri hatırlar."
    )
    
    # 3. Ajanlar
    add_heading(doc, '3. Ajanlar Ekosistemi (Agents Ecosystem)', 1)
    agents = [
        ("PlannerAgent", "Hedefi analiz eder ve adım adım görevlere böler."),
        ("CoderAgent", "Gelişmiş kod yazımı ve modifikasyon yapar. Hızlı mod (CoderFast) desteği vardır."),
        ("ExecutorAgent", "Yerel veya Docker üzerinde terminal komutlarını güvenlik sınırları dahilinde çalıştırır."),
        ("CriticAgent", "Yapılan işi denetler, kod kalitesini (UI ve backend) değerlendirip 10 üzerinden puanlar."),
        ("ResearcherAgent", "İnternette ve proje içinde (RAG ile) detaylı arama yapar."),
        ("BuilderAgent", "Flet, Android APK gibi platform spesifik derleme işlemlerini ve paketlemeyi üstlenir."),
        ("Tester / Linter / UITester", "Sırasıyla Pytest, Pylint ve Playwright tabanlı testleri yürüten kontrolcü ajanlardır.")
    ]
    for agent, desc in agents:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(agent + ": ").bold = True
        p.add_run(desc)
        
    # 4. Araçlar ve Entegrasyon
    add_heading(doc, '4. Araçlar (Tools) ve Servisler', 1)
    doc.add_paragraph(
        "Sistem, OpenAI 'function calling' standardı ile uyumlu '@register_tool' dekoratörünü kullanır. "
        "Bağlı araçlar arasında şunlar yer alır:"
    )
    tools = [
        "Code Runner & Shell Executor",
        "File Manager & Project Indexer",
        "Web Search (DuckDuckGo) & Simple Search (RAG)",
        "Docker Runner & Requirements Generator"
    ]
    for t in tools:
        doc.add_paragraph(t, style='List Bullet')
        
    add_heading(doc, '5. Kullanıcı Arayüzü ve API (UI & Web)', 1)
    doc.add_paragraph(
        "Sistemin iki ana kullanım yöntemi vardır:\n"
        "1. Etkileşimli CLI (Terminal Arayüzü): Gerçek zamanlı log izleme ve manuel kontrol.\n"
        "2. FastAPI Dashboard (main_api.py): Modern, şifre korumalı, asenkron (SSE / Server-Sent Events) "
        "log akışı sağlayan ve mobil cihazlardan da erişilebilen web paneli."
    )
    
    # Kaydetme
    save_path = "MAOS_Kapsamli_Analiz_Raporu.docx"
    doc.save(save_path)
    print(f"Rapor başarıyla oluşturuldu: {save_path}")

if __name__ == '__main__':
    main()
